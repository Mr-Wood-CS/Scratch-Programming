#!/usr/bin/env python3
import html
import sys
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NS = {"q": "http://www.imsglobal.org/xsd/ims_qtiasiv1p2"}


def clean(value):
    value = html.unescape(value or "")
    return " ".join(value.replace("\u00a0", " ").split())


def parse_qti(path):
    root = ET.parse(path).getroot()
    result = []
    for item in root.findall(".//q:item", NS):
        qtype = ""
        for field in item.findall("./q:itemmetadata/q:qtimetadata/q:qtimetadatafield", NS):
            if field.findtext("q:fieldlabel", default="", namespaces=NS) == "question_type":
                qtype = field.findtext("q:fieldentry", default="", namespaces=NS)
        prompt = clean(item.findtext("./q:presentation/q:material/q:mattext", default="", namespaces=NS))
        options = [
            clean(label.findtext("./q:material/q:mattext", default="", namespaces=NS))
            for label in item.findall(".//q:response_label", NS)
        ]
        result.append({"type": qtype, "prompt": prompt, "options": options})
    return result


REPLACEMENTS = {
    "Untitled Question": (
        "Which of these is an example of computer software?",
        ["A keyboard", "Microsoft Word", "A monitor", "A printer"],
        "B",
    ),
    "Saving Files": (
        "Which command stores the latest changes made to an existing file?",
        ["Print", "Save", "Close", "Undo"],
        "B",
    ),
    "Computer Programming": (
        "What is a computer program?",
        [
            "A collection of instructions that a computer follows",
            "A physical part inside a computer",
            "A folder used to store documents",
            "A cable that connects a computer to a network",
        ],
        "A",
    ),
    "Software": (
        "What is software?",
        [
            "The physical parts of a computer",
            "Programs and applications that run on a computer",
            "A cable used to connect devices",
            "Information printed on paper",
        ],
        "B",
    ),
    "The Internet": (
        "Which statement best describes the Internet?",
        [
            "A single website owned by one company",
            "A worldwide network that connects computers and devices",
            "A program used only for writing documents",
            "A folder where downloaded files are stored",
        ],
        "B",
    ),
    "Which Image represents 'Centre text'": (
        "Which formatting command places text midway between the left and right margins?",
        ["Align left", "Centre", "Align right", "Justify"],
        "B",
    ),
    "Which Image represents 'Bold'": (
        "Which formatting command makes selected text appear thicker and darker?",
        ["Italic", "Underline", "Bold", "Highlight"],
        "C",
    ),
    "Which Image represents 'Copy'": (
        "Which command places a duplicate of selected content on the clipboard?",
        ["Cut", "Paste", "Delete", "Copy"],
        "D",
    ),
    "Which Image represents 'Spell Check'": (
        "Which tool identifies words that may be spelled incorrectly?",
        ["Word count", "Find and replace", "Spell check", "Print preview"],
        "C",
    ),
}


FOURTH_OPTIONS = {
    "Which of these software packages is a “spreadsheet”?": "Access",
    "Which of these software packages would you use to write a letter?": "Access",
    "Which of these software packages would you use to create a presentation?": "Access",
    "What is a network?": "A program used to create presentations",
    "What does “e-mail” stand for?": "Express mail",
    "What is an attachment?": "A password used to open an account",
    "Should you meet somebody you have met on-line without telling your parents?": "Only if they send you a photograph first",
    "Which of these is classed as “social media”?": "Microsoft Word",
    "What is a “computer virus”?": "A tool used to organise files",
    "What is an “algorithm”?": "A device used to connect to the Internet",
}


KNOWN_ANSWERS = {
    "Which of these software packages is a “spreadsheet”?": "Excel",
    "Which of these software packages would you use to write a letter?": "Word",
    "Which of these software packages would you use to create a presentation?": "PowerPoint",
    "What is a network?": "A collection of computers linked together",
    "What does “e-mail” stand for?": "Electronic mail",
    "What is an attachment?": "A file that is sent with an e-mail",
    "Should you meet somebody you have met on-line without telling your parents?": "No",
    "Which of these is classed as “social media”?": "Facebook",
    "What is a “computer virus”?": "Software that can damage your files or computer",
    "What is an “algorithm”?": "A set of instructions",
    "What is the difference between “Save” and “Save As”?": "“Save” updates the file, “Save As” creates a new copy with a new name/location",
    "Why is it important to save your work with a sensible name?": "To make the file easier to find later",
    "If you have already opened Word, how would you open a document you have saved?": "Click “File” then “Open” and choose the document",
    "What is a folder and why are they used?": "A folder is a way of storing and organising multiple files together on a computer",
    "What is the Internet?": "A worldwide network of connected computers that share information",
    "How would you find and save an image using Google?": "Open Google Images, search, right click the picture, then choose “Save As”",
    "Should you trust everything you find on the internet?": "No, you should check if the source is reliable",
}


def transform(questions):
    transformed = []
    program_occurrence = 0
    for q in questions:
        prompt = q["prompt"]
        if prompt in REPLACEMENTS:
            new_prompt, options, correct = REPLACEMENTS[prompt]
        elif prompt == "What will this program do?":
            program_occurrence += 1
            if program_occurrence % 2 == 1:
                new_prompt = "A Scratch script starts when the space key is pressed, then repeatedly moves 10 steps and bounces at the edge. What will the sprite do?"
                options = q["options"]
                correct = "B"
            else:
                new_prompt = "A Scratch script repeatedly checks whether the sprite is touching the mouse pointer. When it is, the sprite hides, waits one second, and then shows. What will happen?"
                options = q["options"]
                correct = "A"
        elif prompt == "What will this program draw?":
            new_prompt = "A Scratch script puts the pen down, then repeats four times: move 50 steps and turn 90 degrees. What will it draw?"
            options = q["options"]
            correct = "C"
        else:
            new_prompt = prompt
            options = list(q["options"])
            if new_prompt == "If you have already opened Word, how would you open a document you have saved?":
                options[2] = "Click “Review” then “Spelling and Grammar”"
            if len(options) < 4:
                options.append(FOURTH_OPTIONS[new_prompt])
            correct_text = KNOWN_ANSWERS[new_prompt]
            correct = chr(65 + options.index(correct_text))
        if len(options) != 4 or any(not option for option in options):
            raise ValueError(f"Question does not have four complete options: {new_prompt!r}")
        transformed.append({"prompt": new_prompt, "options": options, "correct": correct})
    return transformed


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_paragraph(paragraph, with_next=False):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext" if with_next else "w:keepLines")
    p_pr.append(keep)


def build_doc(questions, output):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    for name in ("Title", "Heading 1", "Heading 2"):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 2" else "Aptos"
        style.font.color.rgb = RGBColor(0, 0, 0)
    styles["Title"].font.size = Pt(25)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True

    if "Question" not in styles:
        q_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    else:
        q_style = styles["Question"]
    q_style.base_style = styles["Normal"]
    q_style.font.name = "Aptos"
    q_style.font.size = Pt(11)
    q_style.font.bold = True
    q_style.font.color.rgb = RGBColor(0, 0, 0)
    q_style.paragraph_format.space_before = Pt(7)
    q_style.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph("Computer Skills Multiple Choice Question Bank", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro = doc.add_paragraph(
        "Choose one answer for each question. Every question has four options. The answer key begins on a new page after Question 57."
    )
    intro.paragraph_format.space_after = Pt(9)

    for number, question in enumerate(questions, 1):
        if number in {11, 23, 36, 48}:
            doc.add_page_break()
        p = doc.add_paragraph(style="Question")
        p.add_run(f"{number}. {question['prompt']}")
        keep_paragraph(p, with_next=True)
        for index, option in enumerate(question["options"]):
            option_p = doc.add_paragraph(style="Normal")
            option_p.paragraph_format.left_indent = Inches(0.22)
            option_p.paragraph_format.first_line_indent = Inches(-0.02)
            option_p.paragraph_format.space_after = Pt(1)
            option_p.add_run(f"☐  {chr(65 + index)}. {option}")
            keep_paragraph(option_p, with_next=index < 3)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

    answer_heading = doc.add_paragraph("Answer Key", style="Heading 1")
    answer_heading.paragraph_format.page_break_before = True
    keep_paragraph(answer_heading, with_next=True)
    key_intro = doc.add_paragraph("Use this key when marking the question bank.")
    key_intro.paragraph_format.space_after = Pt(8)
    rows = (len(questions) + 3) // 4 + 1
    table = doc.add_table(rows=rows, cols=8)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(0.45), Inches(1.0)] * 4
    headers = ["No", "Answer"] * 4
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        cell.text = headers[i]
        set_cell_margins(cell)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E78")
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    for i, question in enumerate(questions):
        group = i // (rows - 1)
        row_index = i % (rows - 1) + 1
        col = group * 2
        table.cell(row_index, col).text = str(i + 1)
        table.cell(row_index, col + 1).text = question["correct"]
    for row_index, row in enumerate(table.rows[1:], 1):
        for col_index, cell in enumerate(row.cells):
            cell.width = widths[col_index]
            set_cell_margins(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_index % 2 == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "EAF2F8")
                cell._tc.get_or_add_tcPr().append(shading)
    doc.core_properties.title = "Computer Skills Multiple Choice Question Bank"
    doc.core_properties.subject = "Multiple choice computer skills questions"
    doc.save(output)


if __name__ == "__main__":
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    build_doc(transform(parse_qti(source)), output)
